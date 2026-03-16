"""
docx_parser.py — Extract raw text and structure from DOCX files.

Uses python-docx for paragraph-level inspection and mammoth for HTML conversion.
Returns a dict with raw text, mammoth HTML, detected metadata fields, and any
"Additional Information for Staging" section parsed into structured data.
"""

import re
import mammoth
from docx import Document


# Regex patterns for labeled metadata lines
LABEL_PATTERNS = {
    'title':           r'^title\s*:\s*(.+)$',
    'subtitle':        r'^subtitle\s*:\s*(.+)$',
    'author_name':     r'^author(?:\s+name)?\s*:\s*(.+)$',
    'author_title':    r'^author\s+title\s*:\s*(.+)$',
    'topic_tags':      r'^(?:topic\s+)?tag(?:\(s\)|s)?\s*:\s*(.+)$',
    'power_keywords':  r'^power\s+keywords?\s*:\s*(.+)$',
    'photo_credit':    r'^PC\s*:\s*(.+)$',
    'age_groups':      r'^age\s+group(?:\(s\)|s)?\s*:\s*(.+)$',
    'original_url':    r'^original\s*:\s*(.+)$',
}

# Detects the staging instructions heading
_STAGING_HEADING_RE = re.compile(
    r'additional\s+information\s+for\s+staging|staging\s+instructions?',
    re.I,
)


def parse_docx(file_path: str) -> dict:
    """
    Parse a DOCX file and return extracted content.

    Returns:
        {
            raw_text: str,
            mammoth_html: str,
            graph_count: int,
            detected_title: str,
            detected_subtitle: str,
            detected_author_name: str,
            detected_author_title: str,
            detected_topic_tags: list[str],
            detected_power_keywords: list[str],
            staging_instructions: {
                featured_image_url: str,
                featured_image_alt: str,
                related_articles: [{article_url, image_url, tagline}],
                graphs: [{url, label}],
            },
        }
    """
    doc = Document(file_path)

    # Find the staging instructions heading (must be a Heading-style paragraph)
    staging_start = None
    for i, para in enumerate(doc.paragraphs):
        if (_STAGING_HEADING_RE.search(para.text)
                and para.style.name.lower().startswith('heading')):
            staging_start = i
            break

    # Article content is everything before the staging section
    article_paras = (
        doc.paragraphs[:staging_start]
        if staging_start is not None
        else doc.paragraphs
    )

    # Raw text for Claude — article content only, staging section excluded
    raw_text = '\n'.join(p.text for p in article_paras)

    # Mammoth converts the full DOCX; we strip the staging section afterwards
    with open(file_path, 'rb') as f:
        mammoth_result = mammoth.convert_to_html(f)
    mammoth_html = mammoth_result.value

    if staging_start is not None:
        staging_heading = doc.paragraphs[staging_start].text
        mammoth_html = _strip_staging_from_html(mammoth_html, staging_heading)

    # Replace base64 embedded images with [[GRAPH_N]] placeholders
    mammoth_html, graph_count = _extract_graph_placeholders(mammoth_html)

    # Scan article paragraphs for labeled metadata fields
    detected = {
        'title': '', 'subtitle': '', 'author_name': '',
        'author_title': '', 'topic_tags': [], 'power_keywords': [],
        'photo_credit': '', 'age_groups': [], 'original_url': '',
    }
    for para in article_paras:
        text = para.text.strip()
        if not text:
            continue
        for field, pattern in LABEL_PATTERNS.items():
            m = re.match(pattern, text, re.IGNORECASE)
            if m:
                value = m.group(1).strip()
                if field in ('topic_tags', 'power_keywords', 'age_groups'):
                    detected[field] = [t.strip() for t in value.split(',') if t.strip()]
                else:
                    detected[field] = value
                break

    # Parse the staging instructions section if present
    staging_instructions = (
        _parse_staging_instructions(doc.paragraphs[staging_start:])
        if staging_start is not None
        else {}
    )

    return {
        'raw_text': raw_text,
        'mammoth_html': mammoth_html,
        'graph_count': graph_count,
        'detected_title': detected['title'],
        'detected_subtitle': detected['subtitle'],
        'detected_author_name': detected['author_name'],
        'detected_author_title': detected['author_title'],
        'detected_topic_tags': detected['topic_tags'],
        'detected_power_keywords': detected['power_keywords'],
        'detected_photo_credit': detected['photo_credit'],
        'detected_age_groups': detected['age_groups'],
        'detected_original_url': detected['original_url'],
        'staging_instructions': staging_instructions,
    }


def parse_digest_docx(file_path: str) -> dict:
    """
    Parse a fertility digest DOCX (exported from Google Doc).

    Expected document structure:
      Subject line: <email title>
      Preheader: <preheader>   (skipped)
      <intro paragraph(s)>
      <Article title>          (bold run)
      <Article description>
      BUTTON: Read more        (paragraph with hyperlink — any hyperlink text)
      ... (repeat for up to 5 articles)

    Returns:
        {
            title:      str,
            intro_text: str,
            articles:   [{title, url, description}] × up to 5
        }
    """
    from docx.oxml.ns import qn as _qn

    doc = Document(file_path)

    email_title = ''
    intro_lines = []
    articles = []
    current_article = None
    seen_first_bold = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # "Subject line: …" → email title
        m = re.match(r'^Subject\s+line\s*:\s*(.+)$', text, re.I)
        if m:
            email_title = m.group(1).strip()
            continue

        # "Preheader: …" → skip
        if re.match(r'^Preheader\s*:', text, re.I):
            continue

        # "BUTTON: …" → CTA paragraph; extract hyperlink URL
        if re.match(r'^BUTTON\s*:', text, re.I):
            url = _get_hyperlink_url(para, doc, _qn)
            if current_article is not None:
                current_article['url'] = url
                articles.append(current_article)
                current_article = None
            continue

        # Bold paragraph → new article title
        if _is_bold_para(para):
            seen_first_bold = True
            current_article = {'title': text, 'description': '', 'url': ''}
            continue

        # Plain text
        if current_article is not None:
            # First plain line after the title is the description
            if not current_article['description']:
                current_article['description'] = text
        elif not seen_first_bold:
            # We're before the first article — collect as intro
            intro_lines.append(text)

    # Flush any incomplete trailing article
    if current_article is not None:
        articles.append(current_article)

    return {
        'title':      email_title,
        'intro_text': ' '.join(intro_lines),
        'articles':   articles[:5],
    }


def parse_paid_digest_docx(file_path: str) -> dict:
    """
    Parse a paid digest DOCX (exported from Google Doc).

    Supports two formats:

    Format A — section name and URL on the same line:
      Popular this week 1: https://parentdata.org/…
      Popular this week 2: https://parentdata.org/…
      Pregnancy: https://parentdata.org/…

    Format B — section name and URL on separate lines:
      Popular this week
      https://parentdata.org/…
      Pregnancy
      https://parentdata.org/…

    Headings like "Popular this week 1" and "Popular this week 2" are
    normalised into a single section with 2 articles.

    Returns:
        {
            sections: [
                {name: str, articles: [{url: str}]},
                ...
            ]
        }
    """
    doc = Document(file_path)

    sections = []
    url_re = re.compile(r'https?://(?:www\.)?parentdata\.org/\S+', re.I)
    # Strip trailing numbers from headings: "Popular this week 1" → "Popular this week"
    trailing_num_re = re.compile(r'\s+\d+\s*$')

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        url_match = url_re.search(text)

        if url_match:
            url = url_match.group(0).rstrip('.,;:')

            # Check if there's section name text before the URL on the same line
            prefix = text[:url_match.start()].strip().rstrip(':').strip()
            if prefix:
                # "Section name 1: https://…" — extract both section name and URL
                name = trailing_num_re.sub('', prefix).strip()
                if name:
                    if not sections or sections[-1]['name'].lower() != name.lower():
                        sections.append({'name': name, 'articles': []})

            # Append the URL to the current section
            if sections:
                sections[-1]['articles'].append({'url': url})
        else:
            # Pure text line — treat as a section heading
            name = trailing_num_re.sub('', text).strip()
            if not name:
                continue

            # Merge into existing section if name matches the last one
            if sections and sections[-1]['name'].lower() == name.lower():
                continue

            sections.append({'name': name, 'articles': []})

    return {'sections': sections}


def parse_toddler_article_docx(file_path: str) -> dict:
    """
    Parse a ToddlerData article DOCX (exported from Google Doc).

    Expected document structure:
      Subject Line: …
      Preheader: ToddlerData, 18 Months Old
      From Name: …                              (skipped)
      https://parentdata.org/some-article/       (bare URL — only PD link in the doc)
      DISCUSSION QUESTIONS                       (bold heading)
      Here are a few questions…                  (italic intro)
      Does my child meet the CDC milestones…
      What percentile is my child in…
      Do I have any concerns…

    Returns:
        {
            months_old:            str,
            article_url:           str,
            discussion_intro:      str,
            discussion_questions:  [str]
        }
    """
    from docx.oxml.ns import qn as _qn

    doc = Document(file_path)

    months_old = ''
    article_url = ''
    discussion_intro = ''
    discussion_questions = []
    in_questions = False

    _pd_url_re = re.compile(r'https?://(?:www\.)?parentdata\.org/\S+', re.I)

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # "Subject line: …" → skip
        if re.match(r'^Subject\s+line\s*:', text, re.I):
            continue

        # "Preheader: …" → extract months
        m = re.match(r'^Preheader\s*:\s*(.+)$', text, re.I)
        if m:
            months_match = re.search(r'(\d+)\s*months?', m.group(1), re.I)
            if months_match:
                months_old = months_match.group(1)
            continue

        # "From Name: …" → skip
        if re.match(r'^From\s+Name\s*:', text, re.I):
            continue

        # "Discussion Questions" heading (bold, with or without colon)
        if re.match(r'^Discussion\s+Questions?\s*:?\s*$', text, re.I):
            in_questions = True
            continue

        # Lines after the Discussion Questions heading
        if in_questions:
            # First italic paragraph → intro text, not a question
            if not discussion_intro and _is_italic_para(para):
                discussion_intro = text
                continue
            # Numbered question → strip the number prefix
            q_match = re.match(r'^\d+\.\s*(.+)$', text)
            if q_match:
                discussion_questions.append(q_match.group(1).strip())
            elif text:
                # Non-numbered question
                discussion_questions.append(text)
            continue

        # Skip placeholder lines like "[Insert full draft]"
        if re.match(r'^\[.*\]$', text):
            continue

        # ParentData URL — found in text or as a hyperlink in the paragraph
        if not article_url:
            url_match = _pd_url_re.search(text)
            if url_match:
                article_url = url_match.group(0).rstrip('.,;:')
            else:
                # Check hyperlinks embedded in the paragraph XML
                hl_url = _get_hyperlink_url(para, doc, _qn)
                if hl_url and _pd_url_re.match(hl_url):
                    article_url = hl_url

    return {
        'months_old': months_old,
        'article_url': article_url,
        'discussion_intro': discussion_intro,
        'discussion_questions': discussion_questions,
    }


def parse_toddler_qa_docx(file_path: str) -> dict:
    """
    Parse a ToddlerData Q&A DOCX (exported from Google Doc).

    Expected document structure:
      Subject Line: …
      Preheader: ToddlerData, 18 months
      From Name: …                         (skipped)
      <intro paragraph(s)>
      <article URL 1>                      (bare URL or hyperlink)
      <article URL 2>
      <article URL 3>                      (optional)

    Returns:
        {
            months_old:   str,
            intro_text:   str,
            article_urls: [str] × 2–3,
        }
    """
    from docx.oxml.ns import qn as _qn

    doc = Document(file_path)

    _pd_url_re = re.compile(r'https?://(?:www\.)?parentdata\.org/\S+', re.I)

    months_old = ''
    intro_lines = []
    article_urls = []
    seen_first_url = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # "Subject line: …" → skip
        if re.match(r'^Subject\s+line\s*:', text, re.I):
            continue

        # "Preheader: …" → extract months
        m = re.match(r'^Preheader\s*:\s*(.+)$', text, re.I)
        if m:
            months_match = re.search(r'(\d+)\s*months?', m.group(1), re.I)
            if months_match:
                months_old = months_match.group(1)
            continue

        # "From Name: …" → skip
        if re.match(r'^From\s+Name\s*:', text, re.I):
            continue

        # ParentData URL → article link
        url = ''
        url_match = _pd_url_re.search(text)
        if url_match:
            url = url_match.group(0).rstrip('.,;:')
        else:
            hl_url = _get_hyperlink_url(para, doc, _qn)
            if hl_url and _pd_url_re.match(hl_url):
                url = hl_url

        if url:
            seen_first_url = True
            article_urls.append(url)
            continue

        # Plain text before first URL → intro
        if not seen_first_url:
            intro_lines.append(text)

    return {
        'months_old': months_old,
        'intro_text': ' '.join(intro_lines),
        'article_urls': article_urls[:3],
    }


def parse_toddler_digest_docx(file_path: str) -> dict:
    """
    Parse a ToddlerData digest DOCX (exported from Google Doc).

    Expected document structure:
      Subject Line: <email title>
      Preheader: ToddlerData, 18 Months Old
      From Name: …                            (skipped)
      <intro paragraph(s)>
      <article URL>                            (bare URL or hyperlink)
      Title: <card title>                      (bold; "Title:" prefix optional)
      Text: <card description>                 ("Text:" prefix optional)
      … (repeat for up to 3 articles)
      Win of the Month / Win of the Week       (bold heading)
      <quote text>
      —Attribution

    Returns:
        {
            title:            str,
            months_old:       str,
            intro_text:       str,
            articles:         [{title, description, url}] × up to 3,
            win_text:         str,
            win_attribution:  str,
        }
    """
    from docx.oxml.ns import qn as _qn

    doc = Document(file_path)

    _pd_url_re = re.compile(r'https?://(?:www\.)?parentdata\.org/\S+', re.I)

    email_title = ''
    months_old = ''
    intro_lines = []
    articles = []
    win_text = ''
    win_attribution = ''
    in_win = False
    seen_first_url = False
    last_article_needs = None  # 'title', 'description', or None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # "Subject line: …" → email title
        m = re.match(r'^Subject\s+line\s*:\s*(.+)$', text, re.I)
        if m:
            email_title = m.group(1).strip()
            continue

        # "Preheader: …" → extract months
        m = re.match(r'^Preheader\s*:\s*(.+)$', text, re.I)
        if m:
            months_match = re.search(r'(\d+)\s*months?\s*old', m.group(1), re.I)
            if months_match:
                months_old = months_match.group(1)
            continue

        # "From Name: …" → skip
        if re.match(r'^From\s+Name\s*:', text, re.I):
            continue

        # "Win of the Week/Month" heading (bold)
        if re.match(r'^Win\s+of\s+the\s+(Week|Month)\s*:?\s*$', text, re.I):
            in_win = True
            last_article_needs = None
            continue

        # Inside Win section
        if in_win:
            # Attribution line starts with em-dash or hyphen
            if re.match(r'^[\u2014\u2013\-]\s*', text):
                win_attribution = re.sub(r'^[\u2014\u2013\-]\s*', '', text).strip()
            elif not win_text:
                win_text = text
            continue

        # ParentData URL → new article (check text and hyperlinks)
        url = ''
        url_match = _pd_url_re.search(text)
        if url_match:
            url = url_match.group(0).rstrip('.,;:')
        else:
            hl_url = _get_hyperlink_url(para, doc, _qn)
            if hl_url and _pd_url_re.match(hl_url):
                url = hl_url

        if url:
            seen_first_url = True
            # If the URL is embedded in a bold paragraph, the paragraph
            # text IS the title.  Otherwise expect title on next line.
            if _is_bold_para(para) and not url_match:
                articles.append({'title': text, 'description': '', 'url': url})
                last_article_needs = 'description'
            else:
                articles.append({'title': '', 'description': '', 'url': url})
                last_article_needs = 'title'
            continue

        # Fill in title or description for the most recent article
        if last_article_needs and articles:
            # Strip "Title:" or "Text:" labels if present
            value = re.sub(r'^(Title|Text)\s*:\s*', '', text, flags=re.I)
            if last_article_needs == 'title':
                articles[-1]['title'] = value
                last_article_needs = 'description'
            elif last_article_needs == 'description':
                articles[-1]['description'] = value
                last_article_needs = None
            continue

        # Plain text before first article URL → intro
        if not seen_first_url:
            intro_lines.append(text)

    return {
        'title': email_title,
        'months_old': months_old,
        'intro_text': ' '.join(intro_lines),
        'articles': articles[:3],
        'win_text': win_text,
        'win_attribution': win_attribution,
    }


def _get_hyperlink_url(para, doc, qn) -> str:
    """Return the URL of the first hyperlink in a paragraph, or ''."""
    for hl in para._element.findall('.//' + qn('w:hyperlink')):
        r_id = hl.get(qn('r:id'), '')
        if r_id and r_id in doc.part.rels:
            rel = doc.part.rels[r_id]
            target = getattr(rel, '_target', None)
            if target and target.startswith('http'):
                return target
    return ''


def _is_bold_para(para) -> bool:
    """Return True if the first non-empty run in the paragraph is bold."""
    for run in para.runs:
        if run.text.strip():
            return bool(run.bold)
    return False


def _is_italic_para(para) -> bool:
    """Return True if the first non-empty run in the paragraph is italic."""
    for run in para.runs:
        if run.text.strip():
            return bool(run.italic)
    return False


def _strip_staging_from_html(html: str, heading_text: str) -> str:
    """
    Remove the staging instructions section from mammoth HTML.

    Finds the heading text, walks back to the opening <h tag, and returns
    everything before it.  This avoids overly greedy regex matching across
    the whole document (mammoth also inserts <a id="…"> anchors inside
    heading elements, so a simple tag-match won't work).
    """
    idx = html.find(heading_text)
    if idx < 0:
        idx = html.lower().find(heading_text.lower())
    if idx < 0:
        return html
    # Walk back from the heading text to the opening <h tag
    h_start = html.rfind('<h', 0, idx)
    return html[:h_start] if h_start >= 0 else html[:idx]


def _parse_staging_instructions(paragraphs) -> dict:
    """
    Parse the 'Additional Information for Staging' section from DOCX paragraphs.

    Recognises these sub-sections (order and presence may vary):

      Featured Image
        Image: <url>        ← URL may be on this line or the next line alone
        Tag:   <alt text>

      Related Reading N:    ← handles typos like "Reaading"
        Link:  <article url>
        Image: <image url>
        Text:  <tagline>

      Graph N:
        Image: <image url>
        Tag:   <alt text>
    """
    result = {
        'featured_image_url': '',
        'featured_image_alt': '',
        'related_articles': [],
        'graphs': [],
        'wp_url': '',
        'fade_from': '',
    }

    # Collect non-empty lines, skipping the heading line itself
    lines = []
    for para in paragraphs:
        text = para.text.strip()
        if text and not _STAGING_HEADING_RE.search(text):
            lines.append(text)

    sections = []          # [(type_str, {key: value}), ...]
    current_type = None
    current_data = {}

    i = 0
    while i < len(lines):
        line = lines[i]

        # ── Top-level labeled fields (detected anywhere in staging section) ─
        if re.match(r'wp\s*(url)?\s*:', line, re.I):
            result['wp_url'] = re.sub(r'^[^:]+:\s*', '', line).strip()

        elif re.match(r'fade\s*from\s*:', line, re.I):
            result['fade_from'] = re.sub(r'^[^:]+:\s*', '', line).strip()

        # ── Section header detection ──────────────────────────────────────
        elif re.match(r'featured\s+image\s*:?\s*$', line, re.I):
            if current_type:
                sections.append((current_type, dict(current_data)))
            current_type = 'featured'
            current_data = {}

        elif re.match(r'related\s+rea+ding\s+\d+\s*:?\s*$', line, re.I):
            if current_type:
                sections.append((current_type, dict(current_data)))
            current_type = 'related'
            current_data = {}

        elif re.match(r'graph\s+\d+\s*:?\s*$', line, re.I):
            if current_type:
                sections.append((current_type, dict(current_data)))
            current_type = 'graph'
            current_data = {}

        # ── Key: value pairs inside a section ────────────────────────────
        elif current_type is not None:
            m = re.match(r'^(\w+)\s*:\s*(.*)', line)
            if m:
                key = m.group(1).lower()
                val = m.group(2).strip()
                # If the value is empty, the next line may be a bare URL
                if (not val
                        and i + 1 < len(lines)
                        and re.match(r'https?://', lines[i + 1])):
                    i += 1
                    val = lines[i]
                current_data[key] = val

        i += 1

    # Flush the final section
    if current_type:
        sections.append((current_type, dict(current_data)))

    # Map parsed sections to the result structure
    for sec_type, data in sections:
        if sec_type == 'featured':
            result['featured_image_url'] = data.get('image', '').strip()
            result['featured_image_alt'] = data.get('tag', '').strip()
        elif sec_type == 'related':
            result['related_articles'].append({
                'article_url': data.get('link', '').strip(),
                'image_url':   data.get('image', '').strip(),
                'tagline':     data.get('text', '').strip(),
            })
        elif sec_type == 'graph':
            result['graphs'].append({
                'url':   data.get('image', '').strip(),
                'label': data.get('tag', '').strip(),
            })

    return result


def parse_simple_docx(file_path: str) -> dict:
    """
    Parse a "simple email" DOCX (exported from Google Doc).

    Expected document structure:
      <pre-button paragraphs>
      BUTTON GOES HERE          (bold marker line)
      <post-button paragraphs>
      Button Information         (heading)
      Text: <button label>
      Link: <button URL>

    Returns:
        {
            pre_button_html:  str,
            post_button_html: str,
            button_text:      str,
            button_url:       str,
        }
    """
    import io

    doc = Document(file_path)

    # Walk paragraphs and split into three zones:
    #   1) before "BUTTON GOES HERE"
    #   2) after "BUTTON GOES HERE" but before "Button Information"
    #   3) "Button Information" section (metadata)
    pre_paras = []
    post_paras = []
    button_text = ''
    button_url = ''

    zone = 'pre'  # pre -> post -> info
    for para in doc.paragraphs:
        text = para.text.strip()

        if zone == 'pre':
            if re.match(r'^BUTTON\s+GOES\s+HERE$', text, re.I):
                zone = 'post'
                continue
            pre_paras.append(para)

        elif zone == 'post':
            if re.match(r'^Button\s+Information$', text, re.I):
                zone = 'info'
                continue
            post_paras.append(para)

        elif zone == 'info':
            m_text = re.match(r'^Text\s*:\s*(.+)$', text, re.I)
            m_link = re.match(r'^Link\s*:\s*(.+)$', text, re.I)
            if m_text:
                button_text = m_text.group(1).strip()
            elif m_link:
                button_url = m_link.group(1).strip()

    # Convert the pre/post paragraph ranges to HTML via mammoth.
    # We use the full DOCX mammoth conversion, then extract by paragraph markers.
    with open(file_path, 'rb') as f:
        mammoth_result = mammoth.convert_to_html(f)
    full_html = mammoth_result.value

    # Split the mammoth HTML at the "BUTTON GOES HERE" text
    marker_re = re.compile(r'BUTTON\s+GOES\s+HERE', re.I)
    btn_info_re = re.compile(r'Button\s+Information', re.I)

    # Find the marker in the HTML and split around it
    marker_match = marker_re.search(full_html)
    if marker_match:
        # Find the <p> tag containing the marker
        # Walk backwards to find the opening <p
        start = full_html.rfind('<p', 0, marker_match.start())
        # Walk forwards to find the closing </p>
        end = full_html.find('</p>', marker_match.end())
        if end != -1:
            end += len('</p>')

        pre_html = full_html[:start].strip() if start > 0 else ''
        remaining = full_html[end:].strip() if end > 0 else ''
    else:
        pre_html = full_html
        remaining = ''

    # Split remaining at "Button Information"
    info_match = btn_info_re.search(remaining)
    if info_match:
        info_start = remaining.rfind('<', 0, info_match.start())
        post_html = remaining[:info_start].strip() if info_start > 0 else remaining
    else:
        post_html = remaining

    return {
        'pre_button_html': pre_html,
        'post_button_html': post_html,
        'button_text': button_text,
        'button_url': button_url,
    }


def parse_baby_send_a_docx(file_path: str) -> dict:
    """
    Parse a BabyData "Send A" DOCX (exported from Google Doc) for the
    multi-section weekly newsletter.

    Actual document structure (not marker-based):
      Subject Line: <email title>
      Preheader: BabyData, 1 Week Old
      From Name: …                         (skipped)
      <intro paragraph(s)>
      [BOLD] Here's what other parents are asking …   ← Petey Q&A heading
      <question text>
      <answer text>
      Got a question? Ask Petey …                     (skip)
      BUTTON: Ask Petey                               (hyperlink → petey_cta_url)
      [BOLD] Fact or Fiction: <title>                  ← FoF section
      <answer text>
      BUTTON: Read more (link)
      [BOLD] <Article title>                           ← Article section
      <Article subtitle>
      <description text>
      BUTTON: Read more (link)
      <Video title>                                    ← Video section
      SCREENSHOT                                       (hyperlink → thumbnail)
      BUTTON: Watch now (link)

    Returns:
        {
            title:           str,
            age_text:        str,
            intro_text:      str,
            qa_pairs:        [{'question': str, 'answer': str}],
            petey_cta_url:   str,
            fact_or_fiction:  {'title': str, 'answer': str, 'url': str},
            article_card:    {'title': str, 'subtitle': str,
                              'description': str, 'url': str},
            video_card:      {'title': str, 'url': str, 'thumbnail_url': str},
        }
    """
    from docx.oxml.ns import qn as _qn

    doc = Document(file_path)

    title = ''
    age_text = ''
    intro_lines = []

    # Petey Q&A
    petey_heading = ''
    petey_question = ''
    petey_answer_lines = []
    petey_cta_url = ''

    # Fact or Fiction
    fof_title = ''
    fof_answer_lines = []
    fof_url = ''

    # Article card
    article_title = ''
    article_subtitle = ''
    article_desc_lines = []
    article_url = ''

    # Video
    video_title = ''
    video_url = ''
    video_thumbnail = ''

    # State machine: metadata → intro → petey → fof → article → video
    state = 'metadata'

    def _is_separator(text):
        """Skip decorative separators like single special chars."""
        return len(text) <= 3 and not text[0].isalnum()

    def _all_runs_bold(para):
        """True if every non-empty run in the paragraph is bold."""
        runs = [r for r in para.runs if r.text.strip()]
        return bool(runs) and all(r.bold for r in runs)

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Skip decorative separators
        if _is_separator(text):
            continue

        # === Metadata lines ===
        m = re.match(r'^Subject\s+line\s*:\s*(.+)$', text, re.I)
        if m:
            title = m.group(1).strip()
            continue

        m = re.match(r'^Preheader\s*:\s*(.+)$', text, re.I)
        if m:
            age_match = re.search(
                r'(\d+\s*(?:week|month)s?\s*old)', m.group(1), re.I,
            )
            if age_match:
                age_text = age_match.group(1).strip()
            continue

        if re.match(r'^From\s+Name\s*:', text, re.I):
            continue

        # === "Fact or Fiction:" explicit marker (works from any state) ===
        fof_m = re.match(r'^Fact\s+or\s+Fiction\s*:\s*(.+)$', text, re.I)
        if fof_m:
            fof_title = fof_m.group(1).strip()
            state = 'fof'
            continue

        # Also match when the bold heading includes "Fact or Fiction:"
        if _all_runs_bold(para):
            fof_m2 = re.match(r'^Fact\s+or\s+Fiction\s*:\s*(.+)$', text, re.I)
            if fof_m2:
                fof_title = fof_m2.group(1).strip()
                state = 'fof'
                continue

        # === BUTTON handling ===
        btn_m = re.match(r'^BUTTON\s*:\s*(.+)$', text, re.I)
        if btn_m:
            btn_label = btn_m.group(1).strip().lower()
            url = _get_hyperlink_url(para, doc, _qn)

            if 'petey' in btn_label or 'ask' in btn_label:
                petey_cta_url = url
            elif 'watch' in btn_label:
                video_url = url
                state = 'done'
            elif state == 'fof':
                fof_url = url
                state = 'post_fof'
            elif state == 'article':
                article_url = url
                state = 'post_article'
            continue

        # === SCREENSHOT (video thumbnail link) ===
        if text.upper() == 'SCREENSHOT':
            url = _get_hyperlink_url(para, doc, _qn)
            if url:
                video_thumbnail = url
            state = 'video'
            continue

        # === Bold heading transitions ===
        if _all_runs_bold(para):
            if state in ('metadata', 'intro'):
                # First bold heading → Petey Q&A
                state = 'petey'
                petey_heading = text
                continue
            elif state == 'post_fof':
                # First bold heading after FoF → Article
                state = 'article'
                article_title = text
                continue

        # === Content by current state ===
        if state in ('metadata', 'intro'):
            state = 'intro'
            intro_lines.append(text)
        elif state == 'petey':
            # Skip "Got a question? Ask Petey" type lines
            if re.search(r'ask\s+petey|got\s+a\s+question', text, re.I):
                continue
            if not petey_question:
                petey_question = text
            else:
                petey_answer_lines.append(text)
        elif state == 'fof':
            fof_answer_lines.append(text)
        elif state == 'article':
            if not article_subtitle:
                article_subtitle = text
            else:
                article_desc_lines.append(text)
        elif state in ('post_article', 'video'):
            state = 'video'
            if not video_title:
                video_title = text

    return {
        'title': title,
        'age_text': age_text,
        'intro_text': '\n\n'.join(intro_lines),
        'qa_pairs': [{
            'question': petey_question,
            'answer': ' '.join(petey_answer_lines),
        }] if petey_question else [],
        'petey_cta_url': petey_cta_url,
        'fact_or_fiction': {
            'title': fof_title,
            'answer': ' '.join(fof_answer_lines),
            'url': fof_url,
        },
        'article_card': {
            'title': article_title,
            'subtitle': article_subtitle,
            'description': ' '.join(article_desc_lines),
            'url': article_url,
        },
        'video_card': {
            'title': video_title,
            'url': video_url,
            'thumbnail_url': video_thumbnail,
        },
    }


def parse_baby_send_b_docx(file_path: str) -> dict:
    """
    Parse a BabyData "Send B" DOCX (exported from Google Doc) for the
    3-card digest format.

    Expected document structure:
      Subject Line: <email title>
      Preheader: Plus more reads for week 1, from BabyData
      From Name: …                         (skipped)
      <intro paragraph(s)>
      — (divider)
      <Article title>                       (plain text, first line after divider)
      The bottom line: <description lines>
      BUTTON: Read more (link)              (hyperlink URL)
      — (divider, repeat for 3 articles)
      — (divider)
      Real Talk: <topic>
      <intro line>
      <quote text>

    Returns:
        {
            title:          str,
            age_text:       str,
            intro_text:     str,
            articles:       [{'title': str, 'description': str, 'url': str}]
                            × up to 3,
            real_talk_text: str,
        }
    """
    from docx.oxml.ns import qn as _qn

    doc = Document(file_path)

    title = ''
    age_text = ''
    intro_lines = []
    articles = []
    current_article = None
    real_talk_lines = []
    in_real_talk = False
    seen_first_divider = False

    _divider_re = re.compile(r'^[\u2014\u2013\-—–]{1,3}$')
    _bottom_line_re = re.compile(r'^The\s+bottom\s+line\s*:\s*', re.I)

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Subject line
        m = re.match(r'^Subject\s+line\s*:\s*(.+)$', text, re.I)
        if m:
            title = m.group(1).strip()
            continue

        # Preheader — extract age (handles "week 1" or "1 week old" etc.)
        m = re.match(r'^Preheader\s*:\s*(.+)$', text, re.I)
        if m:
            preheader = m.group(1)
            # Try "X months/weeks old" first
            age_match = re.search(
                r'(\d+\s*(?:week|month)s?\s*old)', preheader, re.I,
            )
            if age_match:
                age_text = age_match.group(1).strip()
            else:
                # Try "week X" or "month X" (convert to "X week(s) old")
                age_match = re.search(
                    r'(week|month)\s+(\d+)', preheader, re.I,
                )
                if age_match:
                    unit = age_match.group(1).lower()
                    num = age_match.group(2)
                    plural = 's' if int(num) != 1 else ''
                    age_text = f'{num} {unit}{plural} old'
            continue

        # From Name — skip
        if re.match(r'^From\s+Name\s*:', text, re.I):
            continue

        # Divider line (em dash, en dash, or hyphen)
        if _divider_re.match(text):
            # Flush any pending article
            if current_article is not None:
                articles.append(current_article)
                current_article = None
            seen_first_divider = True
            continue

        # "Real Talk" heading (with optional subtitle like "Real Talk: Postpartum")
        if re.match(r'^Real\s+Talk', text, re.I):
            in_real_talk = True
            # Flush any pending article
            if current_article is not None:
                articles.append(current_article)
                current_article = None
            continue

        # Collecting Real Talk lines
        if in_real_talk:
            # Check if the paragraph is fully italic (= the reader quote)
            runs = para.runs
            all_italic = runs and all(r.italic for r in runs if r.text.strip())
            real_talk_lines.append(('italic' if all_italic else 'normal', text))
            continue

        # BUTTON: … — extract hyperlink URL
        if re.match(r'^BUTTON\s*:', text, re.I):
            url = _get_hyperlink_url(para, doc, _qn)
            if current_article is not None:
                current_article['url'] = url or ''
                articles.append(current_article)
                current_article = None
            continue

        # After first divider: article sections
        if seen_first_divider:
            # Check if paragraph is a DOCX list item (bullet)
            _numPr = para._element.find(
                './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr',
            )
            is_list_item = _numPr is not None

            # "The bottom line:" text → article description
            bl_match = _bottom_line_re.match(text)
            if bl_match:
                desc_text = text[bl_match.end():].strip()
                if current_article is not None and desc_text:
                    current_article['description_items'].append(desc_text)
                continue

            # If no current article, this is the article title
            if current_article is None:
                current_article = {
                    'title': text, 'description_items': [],
                    'has_bullets': False, 'url': '',
                }
            else:
                # Additional description lines
                current_article['description_items'].append(text)
                if is_list_item:
                    current_article['has_bullets'] = True
            continue

        # Before first divider: intro text
        intro_lines.append(text)

    # Flush any incomplete trailing article
    if current_article is not None:
        articles.append(current_article)

    # Convert internal description_items to output format
    for art in articles:
        items = art.pop('description_items', [])
        has_bullets = art.pop('has_bullets', False)
        if has_bullets:
            art['description_bullets'] = items
            art['description'] = ''
        else:
            art['description'] = ' '.join(items)
            art['description_bullets'] = []

    return {
        'title': title,
        'age_text': age_text,
        'intro_text': ' '.join(intro_lines),
        'articles': articles[:3],
        'real_talk_prompt': ' '.join(
            t for kind, t in real_talk_lines if kind == 'normal'
        ),
        'real_talk_quote': ' '.join(
            t for kind, t in real_talk_lines if kind == 'italic'
        ),
    }


def parse_baby_qa_docx(file_path: str) -> dict:
    """
    Parse a BabyData Q&A DOCX (exported from Google Doc).

    Expected document structure:
      Subject Line: <email title>
      Preheader: BabyData, 3 Months Old
      From Name: …                         (skipped)
      <intro paragraph(s)>                  **FILTER OUT "It's Q&A day for BabyData"**
      <article URL 1>                      (bare URL or hyperlink)
      <article URL 2>

    Returns:
        {
            title:      str,
            age_text:   str,
            intro_text: str,
            article_urls: [str] × 2,
        }
    """
    from docx.oxml.ns import qn as _qn

    doc = Document(file_path)

    _pd_url_re = re.compile(r'https?://(?:www\.)?parentdata\.org/\S+', re.I)
    _filter_re = re.compile(r"it['\u2019]?s\s+q&?a\s+day\s+for\s+babydata", re.I)

    title = ''
    age_text = ''
    intro_lines = []
    article_urls = []
    seen_first_url = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Subject line
        m = re.match(r'^Subject\s+line\s*:\s*(.+)$', text, re.I)
        if m:
            title = m.group(1).strip()
            continue

        # Preheader — extract age
        m = re.match(r'^Preheader\s*:\s*(.+)$', text, re.I)
        if m:
            age_match = re.search(
                r'(\d+\s*(?:week|month)s?\s*old)', m.group(1), re.I,
            )
            if age_match:
                age_text = age_match.group(1).strip()
            continue

        # From Name — skip
        if re.match(r'^From\s+Name\s*:', text, re.I):
            continue

        # ParentData URL → article link
        url = ''
        url_match = _pd_url_re.search(text)
        if url_match:
            url = url_match.group(0).rstrip('.,;:')
        else:
            hl_url = _get_hyperlink_url(para, doc, _qn)
            if hl_url and _pd_url_re.match(hl_url):
                url = hl_url

        if url:
            seen_first_url = True
            article_urls.append(url)
            continue

        # Plain text before first URL → intro (filter out "It's Q&A day" line)
        if not seen_first_url:
            if not _filter_re.search(text):
                intro_lines.append(text)

    return {
        'title': title,
        'age_text': age_text,
        'intro_text': ' '.join(intro_lines),
        'article_urls': article_urls,
    }


def _extract_graph_placeholders(html: str) -> tuple:
    """
    Replace base64-encoded <img> tags in mammoth HTML with [[GRAPH_N]] markers.
    Returns (cleaned_html, count).
    """
    count = [0]

    def replacer(m):
        count[0] += 1
        return f'[[GRAPH_{count[0]}]]'

    cleaned = re.sub(
        r'<img\b[^>]*\bsrc="data:image/[^"]*"[^>]*>',
        replacer,
        html,
        flags=re.IGNORECASE,
    )
    return cleaned, count[0]
