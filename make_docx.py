"""
Convert USER_GUIDE.md to USER_GUIDE.docx using python-docx.
Run once: python make_docx.py
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_horizontal_rule(doc):
    """Insert a thin horizontal line (paragraph border)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_code_style(para):
    """Style a paragraph as monospace code block with light grey shading."""
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.right_indent = Inches(0.3)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    # Grey shading
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    for run in para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)


def apply_inline(para, text):
    """
    Add a run to para, handling **bold** and `code` inline markup.
    """
    # Pattern: **bold**, `code`, or plain text
    pattern = re.compile(r'\*\*(.+?)\*\*|`([^`]+)`')
    pos = 0
    for m in pattern.finditer(text):
        # Text before this match
        if m.start() > pos:
            para.add_run(text[pos:m.start()])
        if m.group(1) is not None:
            run = para.add_run(m.group(1))
            run.bold = True
        else:
            run = para.add_run(m.group(2))
            run.font.name = 'Courier New'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        pos = m.end()
    if pos < len(text):
        para.add_run(text[pos:])


def convert(md_path, docx_path):
    doc = Document()

    # ── Default style tweaks ──────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    for i, name in enumerate(['Heading 1', 'Heading 2', 'Heading 3']):
        h = doc.styles[name]
        h.font.name = 'Calibri'
        h.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)  # dark navy
        h.font.size = Pt([20, 15, 12][i])
        if i > 0:
            h.font.color.rgb = RGBColor(0x2E, 0x6D, 0xA4)

    # ── Parse markdown ────────────────────────────────────────────────────
    lines = Path(md_path).read_text(encoding='utf-8').splitlines()

    in_code = False
    code_lines = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # ── Fenced code block ─────────────────────────────────────────────
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                # Emit collected code lines
                for cl in code_lines:
                    p = doc.add_paragraph()
                    p.add_run(cl if cl else ' ')
                    set_code_style(p)
                in_code = False
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # ── Horizontal rule ───────────────────────────────────────────────
        if re.match(r'^-{3,}$', line.strip()):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── Headings ──────────────────────────────────────────────────────
        m = re.match(r'^(#{1,3})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            doc.add_heading(text, level=level)
            i += 1
            continue

        # ── Bullet list items ─────────────────────────────────────────────
        m = re.match(r'^(\s*)[-*]\s+(.*)', line)
        if m:
            indent = len(m.group(1))
            text = m.group(2).strip()
            style_name = 'List Bullet 2' if indent >= 2 else 'List Bullet'
            p = doc.add_paragraph(style=style_name)
            apply_inline(p, text)
            i += 1
            continue

        # ── Blockquote (> Note: ...) ──────────────────────────────────────
        m = re.match(r'^>\s*(.*)', line)
        if m:
            text = m.group(1).strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            # Light blue-grey shading
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'EAF2FB')
            pPr.append(shd)
            run = p.add_run(text)
            run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x8A)
            run.font.italic = True
            i += 1
            continue

        # ── Empty line ────────────────────────────────────────────────────
        if line.strip() == '':
            i += 1
            continue

        # ── Regular paragraph ─────────────────────────────────────────────
        p = doc.add_paragraph()
        apply_inline(p, line.strip())
        i += 1

    doc.save(docx_path)
    print(f'Saved: {docx_path}')


if __name__ == '__main__':
    here = Path(__file__).parent
    convert(here / 'USER_GUIDE.md', here / 'USER_GUIDE.docx')
